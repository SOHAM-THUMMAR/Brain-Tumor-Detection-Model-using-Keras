import os
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docs.config import GRAPHS_DIR
from docs.styles import (
    add_styled_heading, add_body_p, set_cell_background, set_cell_margins, format_run, set_section_header_footer
)


def build_chapters_7_to_10(doc, ch_sec):
    # CHAPTER 7.0 TESTING AND EVALUATION
    set_section_header_footer(ch_sec, header_title="7.0 Testing and Evaluation")
    add_styled_heading(doc, "7.0 TESTING AND EVALUATION", level=1)

    add_styled_heading(doc, "7.1 System Test Cases Matrix", level=2)
    add_body_p(doc, "Table 7.1 outlines five black-box test scenarios executed on the web application:")

    t_tc = doc.add_table(rows=1, cols=4)
    t_tc.alignment = WD_TABLE_ALIGNMENT.CENTER
    tc_hdr = t_tc.rows[0].cells
    tc_hdr[0].text = "Test Purpose"
    tc_hdr[1].text = "Input Provided"
    tc_hdr[2].text = "Expected Result"
    tc_hdr[3].text = "Actual Outcome"
    for cell in tc_hdr:
        set_cell_background(cell, "003366")
        for p in cell.paragraphs:
            for r in p.runs:
                format_run(r, size_pt=10, bold=True, color_rgb=(255, 255, 255))

    test_cases_data = [
        ("Valid Tumor MRI Scan", "tumor_scan_01.jpg (224x224)", "Render 'Tumor' badge, confidence %, and Grad-CAM heatmap", "PASSED - Tumor detected, 99.4% conf, heatmap generated"),
        ("Valid Healthy MRI Scan", "notumor_scan_02.png", "Render 'No Tumor' green badge & confidence %", "PASSED - No Tumor detected, 98.1% conf"),
        ("Oversized Image File", "large_mri.png (15 MB)", "Block upload, render error.html HTTP 400 'File Too Large'", "PASSED - Rendered error.html 400"),
        ("Invalid File Extension", "document.pdf", "Block upload, render error.html HTTP 400 'Invalid File Type'", "PASSED - Rendered error.html 400"),
        ("Corrupted Image File", "corrupted.jpg (0 bytes)", "Catch exception server-side, log traceback, render error.html 500", "PASSED - Rendered error.html 500 cleanly"),
    ]
    for t_purp, t_in, t_exp, t_act in test_cases_data:
        r_c = t_tc.add_row().cells
        r_c[0].text = t_purp
        r_c[1].text = t_in
        r_c[2].text = t_exp
        r_c[3].text = t_act
        for i, cell in enumerate(r_c):
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            for p in cell.paragraphs:
                for r in p.runs:
                    format_run(r, size_pt=10, bold=(i == 0))

    add_body_p(doc, "Table 7.1: Black-Box Web Application and Inference Test Cases Matrix")

    add_styled_heading(doc, "7.2 Experimental Performance Graphs", level=2)
    add_body_p(doc, "The model's experimental evaluation graphs generated during training and testing are embedded below:")

    graphs_info = [
        ("Training vs Validation Recall.png", "Figure 7.1: Training vs Validation Recall Curve across 30 Epochs"),
        ("Training vs Validation Loss.png", "Figure 7.2: Training vs Validation Loss Curve showing Model Convergence"),
        ("Confusion Matrix.png", "Figure 7.3: Test Dataset Confusion Matrix Heatmap (666 Samples)"),
        ("ROC Curve.png", "Figure 7.4: Receiver Operating Characteristic (ROC) Curve (AUC = 1.000)"),
    ]

    for g_filename, g_caption in graphs_info:
        g_path = os.path.join(GRAPHS_DIR, g_filename)
        if os.path.exists(g_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.paragraph_format.space_before = Pt(12)
            p_img.paragraph_format.space_after = Pt(4)
            p_img.add_run().add_picture(g_path, width=Inches(4.5))

            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_cap.paragraph_format.space_after = Pt(16)
            r_cap = p_cap.add_run(g_caption)
            format_run(r_cap, size_pt=10, italic=True)

    doc.add_page_break()

    # CHAPTER 8.0 SCREENSHOTS AND USER MANUAL
    set_section_header_footer(ch_sec, header_title="8.0 Screenshots and User Manual")
    add_styled_heading(doc, "8.0 SCREENSHOTS AND USER MANUAL", level=1)

    add_styled_heading(doc, "8.1 User Operating Instructions", level=2)
    add_body_p(doc, "1. Launch the Flask Web Application by executing 'python start.py' in terminal.")
    add_body_p(doc, "2. Open web browser and navigate to http://127.0.0.1:5000.")
    add_body_p(doc, "3. Drag and drop an axial brain MRI scan image (PNG or JPG format) into the dropzone.")
    add_body_p(doc, "4. Click 'Analyze MRI Scan'. View the classification badge, confidence %, and Grad-CAM heatmap overlay.")
    add_body_p(doc, "5. Click 'Model Performance' in the navigation bar to inspect global recall curves and ROC plots.")

    add_styled_heading(doc, "8.2 Application Interface Placeholders", level=2)

    shots = [
        "[Insert Screenshot: Web Application Upload Homepage (index.html)]\nFigure 8.1: Homepage featuring interactive dropzone and medical disclaimer",
        "[Insert Screenshot: Inference Result Page with Tumor Badge & Grad-CAM Heatmap (result.html)]\nFigure 8.2: Inference view displaying classification badge, confidence, and heatmap overlay",
        "[Insert Screenshot: Model Stats Dashboard (stats.html)]\nFigure 8.3: Metrics dashboard displaying test set recall, accuracy, and evaluation curves",
    ]
    for shot in shots:
        p_s = doc.add_paragraph()
        p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_s.paragraph_format.space_before = Pt(18)
        p_s.paragraph_format.space_after = Pt(18)
        r_s = p_s.add_run(shot)
        format_run(r_s, size_pt=11, bold=True, color_rgb=(100, 100, 100))

    doc.add_page_break()

    # CHAPTER 9.0 LIMITATIONS AND FUTURE ENHANCEMENTS
    set_section_header_footer(ch_sec, header_title="9.0 Limitations & Future Scope")
    add_styled_heading(doc, "9.0 LIMITATIONS AND FUTURE ENHANCEMENTS", level=1)

    add_styled_heading(doc, "9.1 Real System Limitations", level=2)
    add_body_p(doc, "• Binary Classification Only: The model predicts Tumor vs No Tumor but does not differentiate glioma, meningioma, or pituitary tumors.")
    add_body_p(doc, "• Clinical Validation Requirement: Model relies on dataset distribution and requires clinical trial verification prior to diagnostic reliance.")

    add_styled_heading(doc, "9.2 Future Enhancements", level=2)
    add_body_p(doc, "• Multi-Class Tumor Subtyping (Glioma vs Meningioma vs Pituitary).")
    add_body_p(doc, "• 3D MRI Volume Processing using 3D-CNN / Vision Transformers (ViT).")
    add_body_p(doc, "• Cloud Microservice Deployment on AWS/GCP with authentication and DICOM file support.")

    doc.add_page_break()

    # CHAPTER 10.0 CONCLUSION AND DISCUSSION
    set_section_header_footer(ch_sec, header_title="10.0 Conclusion")
    add_styled_heading(doc, "10.0 CONCLUSION AND DISCUSSION", level=1)
    add_body_p(
        doc,
        "In this project, a computer-aided brain tumor detection and explainability system was successfully developed and evaluated. By leveraging a custom 4-stage Convolutional Neural Network trained with class weighting and decision threshold optimization (0.3 threshold), the model achieved 98% overall test accuracy and 100% tumor recall (zero missed cancer cases) across 666 evaluation scans. Integration of Grad-CAM heatmaps provides transparent visual explanations of feature activations. Finally, deployment via Python Flask yields a responsive diagnostic web tool."
    )

    doc.add_page_break()
