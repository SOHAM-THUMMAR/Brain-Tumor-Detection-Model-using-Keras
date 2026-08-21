from docx.enum.table import WD_TABLE_ALIGNMENT
from docs.styles import (
    add_styled_heading, add_body_p, set_cell_background, set_cell_margins, format_run, set_section_header_footer
)


def build_chapters_1_to_3(doc, ch_sec):
    # =============================================================
    # 1.0 INTRODUCTION
    # =============================================================
    set_section_header_footer(ch_sec, header_title="1.0 Introduction")
    add_styled_heading(doc, "1.0 INTRODUCTION", level=1)

    add_styled_heading(doc, "1.1 Project Summary", level=2)
    add_body_p(
        doc,
        "Brain tumors are abnormal cellular growths inside the cranium that exert pressure on brain tissues, leading to severe neurological deficits, motor dysfunction, and life-threatening complications. In modern clinical oncology, Magnetic Resonance Imaging (MRI) serves as the gold standard non-invasive diagnostic modality for visualizing brain tissue structures. However, manual evaluation of multidimensional MRI slices requires expert neuroradiologists, whose availability is severely limited in low-resource medical facilities. Furthermore, qualitative visual inspection is susceptible to human fatigue and inter-observer diagnostic variation."
    )
    add_body_p(
        doc,
        "To mitigate these challenges, this project delivers an end-to-end computer-aided brain tumor detection and interpretability framework. Built upon a 4-stage custom Convolutional Neural Network (CNN) trained with TensorFlow/Keras and wrapped inside a responsive Python Flask web application, the system automates axial MRI scan preprocessing, binary tumor classification, and spatial tumor localization via Gradient-weighted Class Activation Mapping (Grad-CAM)."
    )

    add_styled_heading(doc, "1.2 Purpose: Goals & Objectives", level=2)
    add_body_p(
        doc,
        "The primary purpose of this project is to bridge deep learning research with clinical utility by creating an accessible, accurate, and explainable web diagnostic interface. The core goals and objectives include:"
    )
    add_body_p(doc, "• Develop a lightweight CNN architecture trained on preprocessed 224x224 RGB MRI scans.")
    add_body_p(doc, "• Prioritize high tumor recall (sensitivity) through class weighting to eliminate false negative predictions.")
    add_body_p(doc, "• Implement Grad-CAM heatmap overlay to explain deep feature representations visually.")
    add_body_p(doc, "• Build a production-ready Python Flask application adhering to strict input validation and web standards.")

    add_styled_heading(doc, "1.3 Scope", level=2)
    add_body_p(
        doc,
        "The project scope encompasses dataset preprocessing, binary classification model training, evaluation graph generation, interpretability heatmaps, and Flask server deployment. The scope is strictly bounded as follows:"
    )
    add_body_p(doc, "• What the system CAN do: Classify axial brain MRI images into 'Tumor' vs 'No Tumor', provide classification confidence percentages, generate Grad-CAM visual heatmaps, and display benchmark dataset statistics.")
    add_body_p(doc, "• What the system CANNOT do: Provide multi-class subtyping (e.g. distinguishing glioma vs meningioma), process 3D DICOM volumes directly without 2D slice extraction, or replace professional radiological diagnosis without clinical trial validation.")

    add_styled_heading(doc, "1.4 Technology and Literature Review", level=2)
    add_body_p(
        doc,
        "Recent advancements in deep convolutional networks have revolutionized computer-vision-based medical imaging. Traditional machine learning techniques relied heavily on hand-crafted features such as Haar wavelets, texture descriptors (GLCM), and support vector machines (SVM). In contrast, end-to-end CNNs automatically extract hierarchical spatial representations ranging from low-level edge features to complex anatomical boundaries. Recent research by Selvaraju et al. (2017) demonstrated that Gradient-weighted Class Activation Mapping (Grad-CAM) provides fine-grained visual explainability for CNN decisions, making deep learning models trustworthy for clinical decision support."
    )

    doc.add_page_break()

    # =============================================================
    # 2.0 PROJECT MANAGEMENT
    # =============================================================
    set_section_header_footer(ch_sec, header_title="2.0 Project Management")
    add_styled_heading(doc, "2.0 PROJECT MANAGEMENT", level=1)

    add_styled_heading(doc, "2.1 Project Planning and Scheduling", level=2)

    add_styled_heading(doc, "2.1.1 Project Development Approach and Justification", level=3)
    add_body_p(
        doc,
        "An Agile iterative development methodology was adopted for this project. Iterative sprints enabled continuous feedback loops across data curation, model selection, loss metric tuning, explainability implementation, and web front-end development."
    )

    add_styled_heading(doc, "2.1.2 Project Plan including Milestones, Deliverables, Roles and Dependencies", level=3)
    add_body_p(doc, "The project execution was structured into five distinct phases over a 16-week timeline as summarized in Table 2.1 below:")

    t_milestones = doc.add_table(rows=1, cols=4)
    t_milestones.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_hdr = t_milestones.rows[0].cells
    m_hdr[0].text = "Phase"
    m_hdr[1].text = "Milestone Task"
    m_hdr[2].text = "Timeline"
    m_hdr[3].text = "Deliverables & Dependencies"
    set_cell_background(m_hdr[0], "003366")
    set_cell_background(m_hdr[1], "003366")
    set_cell_background(m_hdr[2], "003366")
    set_cell_background(m_hdr[3], "003366")
    for cell in m_hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                format_run(r, size_pt=10, bold=True, color_rgb=(255, 255, 255))

    milestones_data = [
        ("Phase 1", "Dataset Acquisition & Preprocessing", "Weeks 1–3", "Split dataset (Train/Val/Test), normalized arrays"),
        ("Phase 2", "CNN Model Architecture & Training", "Weeks 4–7", "bestModel.keras, history metrics, loss curves"),
        ("Phase 3", "Model Evaluation & Recall Tuning", "Weeks 8–10", "Confusion matrix, ROC-AUC, 100% recall threshold"),
        ("Phase 4", "Flask Web Application & Grad-CAM", "Weeks 11–14", "app.py, model_utils.py, HTML/CSS UI, start.py"),
        ("Phase 5", "Testing & Documentation", "Weeks 15–16", "Project report, unit tests, GitHub repo"),
    ]
    for p_id, m_task, t_line, deliv in milestones_data:
        r_c = t_milestones.add_row().cells
        r_c[0].text = p_id
        r_c[1].text = m_task
        r_c[2].text = t_line
        r_c[3].text = deliv
        for i, cell in enumerate(r_c):
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            for p in cell.paragraphs:
                for r in p.runs:
                    format_run(r, size_pt=10, bold=(i == 0))

    add_body_p(doc, "Table 2.1: Project Development Milestones and Execution Schedule")

    add_styled_heading(doc, "2.1.3 Schedule Representation", level=3)
    add_body_p(
        doc,
        "The project schedule was managed using a Gantt timeline representation spanning 16 weeks. Data preparation occupied Weeks 1-3, Model Architecture & Training spanned Weeks 4-7, Recall Tuning & Metric Validation spanned Weeks 8-10, Flask & Grad-CAM Web App Integration spanned Weeks 11-14, and Verification, Report Generation, and Deployment Documentation completed in Weeks 15-16."
    )

    add_styled_heading(doc, "2.2 Risk Management", level=2)

    add_styled_heading(doc, "2.2.1 Risk Identification", level=3)
    add_body_p(
        doc,
        "Risk identification focused on discovering potential technical, medical, and operational risks specific to this deep learning medical diagnostic application. Four primary risks were identified: (1) False negative diagnosis (missed tumor cases), (2) Overfitting to training MRI scans, (3) Corrupted or malicious web file uploads, and (4) Grad-CAM feature map layer extraction failures."
    )

    add_styled_heading(doc, "2.2.2 Risk Analysis", level=3)
    add_body_p(
        doc,
        "Risk analysis evaluated the probability and severity of each identified risk. False negative classification carried CRITICAL severity because missing a malignant brain tumor in screening has severe life-threatening consequences. Overfitting carried HIGH impact as it impairs generalization on unseen patient scans. Malicious file uploads carried MEDIUM impact on server stability, while Grad-CAM layer mismatch carried LOW impact on prediction output."
    )

    add_styled_heading(doc, "2.2.3 Risk Planning", level=3)
    add_body_p(
        doc,
        "Risk planning formulated concrete strategies implemented directly in the code to manage each risk, as detailed in Table 2.2:"
    )

    t_risk = doc.add_table(rows=1, cols=3)
    t_risk.alignment = WD_TABLE_ALIGNMENT.CENTER
    rk_hdr = t_risk.rows[0].cells
    rk_hdr[0].text = "Identified Risk"
    rk_hdr[1].text = "Impact Level"
    rk_hdr[2].text = "Mitigation Strategy Implemented"
    for cell in rk_hdr:
        set_cell_background(cell, "003366")
        for p in cell.paragraphs:
            for r in p.runs:
                format_run(r, size_pt=10, bold=True, color_rgb=(255, 255, 255))

    risk_data = [
        ("False Negative Risk (Missed Cancer)", "CRITICAL", "Assigned 3x class weight to tumor class; set classification threshold to 0.3"),
        ("Model Overfitting", "HIGH", "Added Dropout layers (0.25) and EarlyStopping monitored on val_recall"),
        ("Invalid / Corrupted Web Uploads", "MEDIUM", "Server-side file extension, size limit (10MB), and try-except handling"),
        ("Grad-CAM Layer Failure", "LOW", "Dynamic Conv2D layer inspection with graceful error fallback"),
    ]
    for r_name, r_imp, r_mit in risk_data:
        r_c = t_risk.add_row().cells
        r_c[0].text = r_name
        r_c[1].text = r_imp
        r_c[2].text = r_mit
        for i, cell in enumerate(r_c):
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            for p in cell.paragraphs:
                for r in p.runs:
                    format_run(r, size_pt=10, bold=(i == 0))

    add_body_p(doc, "Table 2.2: Risk Assessment and Mitigation Strategy Matrix")

    doc.add_page_break()

    # =============================================================
    # 3.0 SYSTEM REQUIREMENTS STUDY
    # =============================================================
    set_section_header_footer(ch_sec, header_title="3.0 System Requirements Study")
    add_styled_heading(doc, "3.0 SYSTEM REQUIREMENTS STUDY", level=1)

    add_styled_heading(doc, "3.1 User Characteristics", level=2)
    add_body_p(
        doc,
        "The system is designed for medical researchers, computer science students, and biomedical developers. Users are expected to have basic familiarity with web interfaces and digital brain MRI scan file formats (PNG, JPG, JPEG)."
    )

    add_styled_heading(doc, "3.2 Hardware and Software Requirements", level=2)
    add_body_p(doc, "Table 3.1 delineates the minimum and recommended hardware/software specifications for both model training and web hosting:")

    t_hw = doc.add_table(rows=1, cols=3)
    t_hw.alignment = WD_TABLE_ALIGNMENT.CENTER
    hw_hdr = t_hw.rows[0].cells
    hw_hdr[0].text = "Component"
    hw_hdr[1].text = "Minimum Requirement"
    hw_hdr[2].text = "Recommended Specification"
    for cell in hw_hdr:
        set_cell_background(cell, "003366")
        for p in cell.paragraphs:
            for r in p.runs:
                format_run(r, size_pt=10, bold=True, color_rgb=(255, 255, 255))

    hw_data = [
        ("Processor (CPU)", "Dual-Core 2.0 GHz Intel/AMD", "Quad-Core 3.0 GHz Intel i7 / AMD Ryzen"),
        ("System Memory (RAM)", "8 GB DDR4", "16 GB DDR4 / DDR5"),
        ("GPU Acceleration", "CPU execution supported", "NVIDIA RTX 3060 (CUDA / cuDNN)"),
        ("Disk Storage", "2 GB free space", "10 GB NVMe SSD"),
        ("Operating System", "Windows 10/11, Ubuntu 20.04, macOS", "Windows 11 / Linux Ubuntu 22.04 LTS"),
        ("Python Environment", "Python 3.9+", "Python 3.11.x Virtual Environment"),
        ("Primary Libraries", "TensorFlow 2.12+, Flask 3.0+", "TensorFlow 2.21, Flask 3.1, OpenCV, NumPy"),
    ]
    for comp, req_m, req_r in hw_data:
        r_c = t_hw.add_row().cells
        r_c[0].text = comp
        r_c[1].text = req_m
        r_c[2].text = req_r
        for i, cell in enumerate(r_c):
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            for p in cell.paragraphs:
                for r in p.runs:
                    format_run(r, size_pt=10, bold=(i == 0))

    add_body_p(doc, "Table 3.1: Hardware and Software System Specifications")

    add_styled_heading(doc, "3.3 Constraints", level=2)
    add_body_p(
        doc,
        "System implementation and operation are subject to several technical and regulatory constraints:\n"
        "• Regulatory & Clinical Constraints: The system is developed exclusively for educational and decision-support research. It is not approved by regulatory medical bodies (e.g. FDA/CE) for autonomous diagnostic use.\n"
        "• Hardware Limitations: Inference runs efficiently on standard CPUs without requiring high-end dedicated GPUs.\n"
        "• File Size & Format Constraints: Uploads are constrained to PNG, JPG, and JPEG formats under 10MB per request.\n"
        "• Security & Reliability Constraints: Strict input validation prevents arbitrary file execution, and exception boundaries ensure the web server returns clean HTTP 400/500 responses without exposing internal stack traces."
    )

    doc.add_page_break()
