from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docs.styles import (
    format_run, add_styled_heading, add_body_p,
    set_cell_background, set_cell_margins
)


def build_abbreviations(doc):
    add_styled_heading(doc, "NOTATIONS AND ABBREVIATIONS", level=1)

    table_abbr = doc.add_table(rows=1, cols=2)
    table_abbr.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table_abbr.rows[0].cells
    hdr[0].text = "Abbreviation"
    hdr[1].text = "Full Description"
    set_cell_background(hdr[0], "003366")
    set_cell_background(hdr[1], "003366")
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                format_run(r, size_pt=11, bold=True, color_rgb=(255, 255, 255))

    abbr_data = [
        ("CNN", "Convolutional Neural Network"),
        ("MRI", "Magnetic Resonance Imaging"),
        ("Grad-CAM", "Gradient-weighted Class Activation Mapping"),
        ("ROC", "Receiver Operating Characteristic"),
        ("AUC", "Area Under Curve"),
        ("ReLU", "Rectified Linear Unit"),
        ("API", "Application Programming Interface"),
        ("HTTP", "Hypertext Transfer Protocol"),
        ("UI/UX", "User Interface / User Experience"),
        ("FN", "False Negative (Missed Diagnosis)"),
        ("FP", "False Positive (False Alarm)"),
        ("TN", "True Negative"),
        ("TP", "True Positive"),
        ("PEP", "Python Enhancement Proposal"),
    ]

    for abbr, desc in abbr_data:
        row_cells = table_abbr.add_row().cells
        row_cells[0].text = abbr
        row_cells[1].text = desc
        for i, cell in enumerate(row_cells):
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            for p in cell.paragraphs:
                for r in p.runs:
                    format_run(r, size_pt=11, bold=(i == 0))

    doc.add_page_break()


def build_toc_and_lists(doc):
    add_styled_heading(doc, "LIST OF FIGURES", level=1)
    fig_list = [
        "Figure 4.1: Data Flow Diagram (DFD) of Brain Tumor Classification Pipeline",
        "Figure 4.2: High-Level System Class Architecture Diagram",
        "Figure 5.1: Overall System Architecture of Flask Web App & Keras Model",
        "Figure 5.2: Procedural Logic Flowchart for Image Upload and Inference",
        "Figure 7.1: Training vs Validation Recall Curve across 30 Epochs",
        "Figure 7.2: Training vs Validation Loss Curve showing Model Convergence",
        "Figure 7.3: Test Dataset Confusion Matrix Heatmap (666 Samples)",
        "Figure 7.4: Receiver Operating Characteristic (ROC) Curve (AUC = 1.000)",
        "Figure 8.1: Web Interface Upload Homepage (Drag-and-Drop Dropzone)",
        "Figure 8.2: Inference Result Page with Tumor Badge & Grad-CAM Heatmap",
        "Figure 8.3: Model Metrics & Performance Dashboard Page",
    ]
    for fig in fig_list:
        add_body_p(doc, fig)

    add_styled_heading(doc, "LIST OF TABLES", level=1, space_before=18)
    tab_list = [
        "Table 2.1: Project Development Milestones and Execution Schedule",
        "Table 2.2: Risk Assessment and Mitigation Strategy Matrix",
        "Table 3.1: Hardware and Software System Specifications",
        "Table 4.1: Dataset Distribution across Train, Validation, and Test Sets",
        "Table 7.1: Black-Box Web Application and Inference Test Cases Matrix",
    ]
    for tab in tab_list:
        add_body_p(doc, tab)

    add_styled_heading(doc, "TABLE OF CONTENTS", level=1, space_before=18)
    toc_data = [
        ("CANDIDATE'S DECLARATION", "iii"),
        ("CERTIFICATE", "iv"),
        ("ACKNOWLEDGEMENT", "v"),
        ("ABSTRACT", "vi"),
        ("NOTATIONS AND ABBREVIATIONS", "vii"),
        ("1.0 INTRODUCTION", "1"),
        ("    1.1 Project Summary", "1"),
        ("    1.2 Purpose and Objectives", "2"),
        ("    1.3 Project Scope and Limitations", "3"),
        ("    1.4 Literature Review on CNN Medical Diagnosis", "4"),
        ("2.0 PROJECT MANAGEMENT", "5"),
        ("    2.1 Development Approach", "5"),
        ("    2.2 Project Schedule and Milestones", "6"),
        ("    2.3 Risk Identification and Mitigation Plan", "7"),
        ("3.0 SYSTEM REQUIREMENTS STUDY", "8"),
        ("    3.1 User Characteristics", "8"),
        ("    3.2 Hardware Requirements", "8"),
        ("    3.3 Software Requirements", "9"),
        ("    3.4 Operational and Clinical Constraints", "9"),
        ("4.0 SYSTEM ANALYSIS", "10"),
        ("    4.1 Analysis of Existing System", "10"),
        ("    4.2 Proposed Automated CNN System", "11"),
        ("    4.3 Feasibility Study", "12"),
        ("    4.4 Data Modeling and Dictionary", "13"),
        ("    4.5 Functional Modeling (Data Flow)", "14"),
        ("5.0 SYSTEM DESIGN", "15"),
        ("    5.1 System Architecture Design", "15"),
        ("    5.2 Data Design and Folder Hierarchy", "16"),
        ("    5.3 Procedural Logic Flow", "17"),
        ("    5.4 Input/Output and Interface Design", "18"),
        ("6.0 IMPLEMENTATION PLANNING", "19"),
        ("    6.1 Development Environment Setup", "19"),
        ("    6.2 Module Specifications", "20"),
        ("    6.3 Security Features & Input Validation", "21"),
        ("    6.4 Core Implementation Code Snippet", "22"),
        ("7.0 TESTING AND EVALUATION", "23"),
        ("    7.1 Testing Strategy and Black-Box Methodology", "23"),
        ("    7.2 System Test Cases Matrix", "24"),
        ("    7.3 Experimental Performance Analysis & Figures", "25"),
        ("8.0 SCREENSHOTS AND USER MANUAL", "28"),
        ("    8.1 Web Interface User Manual", "28"),
        ("    8.2 System Application Screenshots", "29"),
        ("9.0 LIMITATIONS AND FUTURE ENHANCEMENTS", "31"),
        ("    9.1 System Limitations", "31"),
        ("    9.2 Future Scope and Enhancements", "32"),
        ("10.0 CONCLUSION AND DISCUSSION", "33"),
        ("REFERENCES", "34"),
        ("APPENDICES", "35"),
    ]

    for title, pg in toc_data:
        p_toc = doc.add_paragraph()
        p_toc.paragraph_format.line_spacing = 1.2
        p_toc.paragraph_format.space_after = Pt(3)
        r_t = p_toc.add_run(f"{title} ".ljust(60, "."))
        format_run(r_t, size_pt=11, bold=title.isupper())
        r_p = p_toc.add_run(f"  {pg}")
        format_run(r_p, size_pt=11, bold=True)

    doc.add_page_break()
