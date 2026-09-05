from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docs.config import (
    STUDENT_NAME, ENROLLMENT_NO, STUDENT_1, STUDENT_2, PROJECT_TITLE, FULL_TITLE,
    DEPARTMENT, GUIDE_NAME, HOD_NAME, INSTITUTE_NAME, UNIVERSITY_NAME, ACADEMIC_YEAR, SUBMISSION_DATE
)
from docs.styles import format_run, add_styled_heading, add_body_p


def build_cover_page(doc):
    p_cov_title = doc.add_paragraph()
    p_cov_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cov_title.paragraph_format.space_before = Pt(36)
    r = p_cov_title.add_run(f"{PROJECT_TITLE}\n\n")
    format_run(r, size_pt=20, bold=True, color_rgb=(0, 51, 102))

    r = p_cov_title.add_run("A PROJECT REPORT\n")
    format_run(r, size_pt=14, bold=True)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(18)
    p_sub.paragraph_format.space_after = Pt(24)
    r = p_sub.add_run("SUBMITTED IN PARTIAL FULFILLMENT OF THE REQUIREMENT FOR THE AWARD OF THE DEGREE OF\n\n")
    format_run(r, size_pt=11, italic=True)
    r = p_sub.add_run(f"B.TECH. ({DEPARTMENT.upper()}) TO\n")
    format_run(r, size_pt=13, bold=True)
    r = p_sub.add_run(f"{UNIVERSITY_NAME.upper()}\n")
    format_run(r, size_pt=13, bold=True)

    p_by = doc.add_paragraph()
    p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_by.paragraph_format.space_before = Pt(24)
    p_by.paragraph_format.space_after = Pt(24)
    r = p_by.add_run("SUBMITTED BY\n\n")
    format_run(r, size_pt=12, bold=True)
    r = p_by.add_run(f"Name of Student\t\tEnrollment No.\n{STUDENT_1}\n{STUDENT_2}\n\n")
    format_run(r, size_pt=12)

    r = p_by.add_run("UNDER THE GUIDANCE OF\n\n")
    format_run(r, size_pt=12, bold=True)
    r = p_by.add_run(f"Internal Guide: {GUIDE_NAME}\n{INSTITUTE_NAME}\n\n")
    format_run(r, size_pt=12)

    r = p_by.add_run(f"{SUBMISSION_DATE}\n\n{INSTITUTE_NAME.upper()}\n")
    format_run(r, size_pt=12, bold=True)

    doc.add_page_break()


def build_title_page(doc):
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_t.paragraph_format.space_before = Pt(36)
    r = p_t.add_run(f"{PROJECT_TITLE}\n\n")
    format_run(r, size_pt=16, bold=True)
    r = p_t.add_run("B.Tech Project Report\n")
    format_run(r, size_pt=14, italic=True)

    add_body_p(doc, f"Project Title: {FULL_TITLE}")
    add_body_p(doc, f"Submitted By: {STUDENT_1}, {STUDENT_2}")
    add_body_p(doc, f"Branch / Department: {DEPARTMENT}")
    add_body_p(doc, f"Internal Guide: {GUIDE_NAME}")
    add_body_p(doc, f"Head of Department: {HOD_NAME}")
    add_body_p(doc, f"Institution: {INSTITUTE_NAME}")
    add_body_p(doc, f"Date of Submission: {SUBMISSION_DATE}")

    doc.add_page_break()


def build_declaration(doc):
    add_styled_heading(doc, "DECLARATION", level=1)
    add_body_p(
        doc,
        f"We hereby certify that We are the sole authors of this project work and that neither any part of this project work nor the whole of the project work has been submitted for a degree to any other University or Institution. We certify that, to the best of our knowledge, our project work does not infringe upon anyone’s copyright nor violate any proprietary rights and that any ideas, techniques, quotations, or any other material from the work of other people included in my/our project document, published or otherwise, are fully acknowledged in accordance with the standard referencing practices. We declare that this is a true copy of our project work, including any final revisions, as approved by my/our project review committee."
    )

    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_before = Pt(48)
    r = p_sig.add_run(
        f"Signature of Student (S)\n\n"
        f"{STUDENT_1}\t\t{STUDENT_2}\n"
        f"Date: {SUBMISSION_DATE}\t\tDate: {SUBMISSION_DATE}\n"
        f"Place: Rajkot, Gujarat\t\tPlace: Rajkot, Gujarat"
    )
    format_run(r, size_pt=11, bold=True)

    doc.add_page_break()


def build_certificate(doc):
    add_styled_heading(doc, "CERTIFICATE", level=1)
    add_body_p(
        doc,
        f"This is to certify that the work which is being presented in the Project Report entitled \"{PROJECT_TITLE}\", in partial fulfillment of the requirement for the award of the degree of B.Tech. (Computer Engineering) and submitted to the {INSTITUTE_NAME}, is an authentic record of our own work carried out during a period from June 2026 to December 2026."
    )
    add_body_p(
        doc,
        "The matter presented in this Project Report has not been submitted by us for the award of any other degree elsewhere."
    )

    p_cert_sig = doc.add_paragraph()
    p_cert_sig.paragraph_format.space_before = Pt(48)
    r = p_cert_sig.add_run(
        f"Signature of Student (S)\n\n"
        f"{STUDENT_1}, {STUDENT_2}\n\n\n"
        f"This is to certify that the above statement made by the students is correct to the best of my knowledge.\n\n\n"
        f"Internal Guide:\t\t\t\tHead of Department:\n"
        f"{GUIDE_NAME}\t\t\t\t{HOD_NAME}\n"
        f"Assistant Professor, CE / IT\t\t\t\tCE / IT, {INSTITUTE_NAME}\n"
        f"RK University, Rajkot\t\t\t\tRK University, Rajkot\n\n"
        f"{SUBMISSION_DATE}"
    )
    format_run(r, size_pt=11, bold=True)

    doc.add_page_break()


def build_acknowledgement(doc):
    add_styled_heading(doc, "ACKNOWLEDGEMENT", level=1)
    add_body_p(
        doc,
        f"We express our sincere gratitude and appreciation to our project guide, {GUIDE_NAME}, for providing valuable guidance, continuous support, and constructive suggestions throughout the development of this project. Her technical expertise and encouragement greatly contributed to the successful implementation of the Brain Tumor Detection System, including the CNN model, model evaluation, Grad-CAM visualization, and web application."
    )
    add_body_p(
        doc,
        f"We would also like to express our sincere thanks to the Head of Department, {HOD_NAME}, for providing the necessary academic guidance, computational resources, and facilities required for carrying out the project."
    )
    add_body_p(
        doc,
        f"We are grateful to all faculty members of the Department of Computer Engineering / IT for their valuable suggestions, technical guidance, and support throughout the project work."
    )
    add_body_p(
        doc,
        "Finally, We extend our sincere appreciation to everyone in the department who provided their valuable assistance and cooperation during the course of this project."
    )

    doc.add_page_break()


def build_abstract(doc):
    add_styled_heading(doc, "ABSTRACT", level=1)
    add_body_p(
        doc,
        "Brain tumors represent one of the most critical and life-threatening neurological conditions. Early detection and precise localization of brain neoplasms from Magnetic Resonance Imaging (MRI) scans are vital for effective surgical planning and treatment outcomes. Manual evaluation of brain MRI slices by radiologists is time-consuming and subject to inter-observer variability. In this work, an end-to-end automated computer-aided diagnosis system is developed using custom Deep Convolutional Neural Networks (CNN) implemented in TensorFlow/Keras, integrated with a clinical-ready Python Flask web interface."
    )
    add_body_p(
        doc,
        "The proposed deep CNN model consists of four sequential convolutional blocks with ReLU activations, max-pooling, dropout regularization, and a sigmoid classification head optimized via binary crossentropy loss. To eliminate false-negative diagnosis in screening scenarios, class-weighted training and decision threshold tuning (0.3 threshold) were employed. On a comprehensive test dataset of 666 MRI slices (458 tumor, 208 non-tumor), the model achieved 98% overall accuracy, 100% tumor recall (zero missed tumor cases), and an Area Under ROC Curve (AUC) of 1.000."
    )
    add_body_p(
        doc,
        "To make the model interpretable and clinically actionable, Gradient-weighted Class Activation Mapping (Grad-CAM) and explicit contour/bounding box spatial localization were integrated to highlight tumor regions. Furthermore, an automated PDF patient diagnostic report microservice (ReportLab) was implemented, enabling one-click download of clinical reports featuring embedded MRI scans, diagnostic badges, confidence percentages, and benchmark model metrics."
    )

    doc.add_page_break()
