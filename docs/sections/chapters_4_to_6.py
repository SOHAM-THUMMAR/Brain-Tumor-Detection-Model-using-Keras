from docx.shared import Pt, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docs.styles import (
    add_styled_heading, add_body_p, set_cell_background, set_cell_margins, format_run, set_section_header_footer
)


def build_chapters_4_to_6(doc, ch_sec):
    # =============================================================
    # 4.0 SYSTEM ANALYSIS
    # =============================================================
    set_section_header_footer(ch_sec, header_title="4.0 System Analysis")
    add_styled_heading(doc, "4.0 SYSTEM ANALYSIS", level=1)

    add_styled_heading(doc, "4.1 Study of Current System", level=2)
    add_body_p(
        doc,
        "The current diagnostic system relies on manual qualitative evaluation of brain MRI slices by neuroradiologists. Radiologists inspect T1-weighted, T2-weighted, and FLAIR MRI sequences to identify abnormal tissue contrast, edema, and mass effect."
    )

    add_styled_heading(doc, "4.2 Problem and Weaknesses of Current System", level=2)
    add_body_p(
        doc,
        "Key weaknesses of manual evaluation include: (1) High diagnostic latency during emergency triaging, (2) Subjective boundary interpretation leading to missed low-contrast or early-stage tumors, (3) High cognitive workload causing diagnostic fatigue, (4) Lack of automated quantitative explainability, and (5) Lack of instant downloadable patient diagnostic reports."
    )

    add_styled_heading(doc, "4.3 Requirements of Proposed System", level=2)
    add_body_p(
        doc,
        "The proposed system addresses these limitations through automated deep learning inference and report generation:\n"
        "• Functional Requirements: Upload axial brain MRI scans via browser, perform 224x224 RGB image preprocessing, classify scans into 'Tumor' or 'No Tumor' with confidence percentage, generate Grad-CAM visual heatmaps, produce explicit tumor region bounding box & contour highlights, generate downloadable PDF Patient Diagnostic Reports, and display historical benchmark model performance metrics.\n"
        "• Non-Functional Requirements: Sub-second inference time (<500ms), 100% tumor recall safety target, robust HTTP 400/500 error views, automatic frontend file selection validation, intuitive responsive dark UI usable on desktop and mobile, and full compliance with PEP8 code modularity."
    )

    add_styled_heading(doc, "4.4 Feasibility Study", level=2)
    add_body_p(
        doc,
        "Feasibility analysis was performed across three dimensions:\n"
        "• Technical Feasibility: High. TensorFlow/Keras, OpenCV, and ReportLab provide mature, robust APIs for deep learning inference, Grad-CAM visualization, and PDF generation on standard hardware.\n"
        "• Operational Feasibility: High. The single-page drag-and-drop web UI requires zero technical training for medical researchers or healthcare students.\n"
        "• Economic Feasibility: High. Built entirely using open-source Python software stack (Flask, Keras, OpenCV, NumPy, ReportLab) without licensing costs."
    )

    add_styled_heading(doc, "4.5 Requirements Validation", level=2)
    add_body_p(
        doc,
        "Requirements were validated through empirical testing on a benchmark dataset of 666 MRI test scans, demonstrating 98% accuracy, zero false negatives (100% recall), and verified PDF report generation."
    )

    add_styled_heading(doc, "4.6 Functions of System", level=2)
    add_styled_heading(doc, "4.6.1 Use Cases and System Scenarios", level=3)
    add_body_p(
        doc,
        "Primary Use Case Scenario:\n"
        "1. User accesses web homepage at http://127.0.0.1:5000 via one-click start.py script.\n"
        "2. User selects or drops a brain MRI image file (PNG/JPG).\n"
        "3. System validates file extension and size.\n"
        "4. Server executes preprocessing, CNN prediction, Grad-CAM heatmap generation, tumor bounding box contour highlighting, and PDF report creation.\n"
        "5. System renders results page displaying color-coded prediction badge, confidence bar, 3-card image comparison grid (Original, Tumor Highlight, Grad-CAM Heatmap), and a 'Download PDF Diagnostic Report' button."
    )

    add_styled_heading(doc, "4.7 Data Modeling", level=2)

    add_styled_heading(doc, "4.7.1 Data Dictionary", level=3)
    add_body_p(
        doc,
        "Data Elements in Image Processing Pipeline:\n"
        "• raw_image_file: Binary image stream uploaded by user (PNG/JPG, max 10MB).\n"
        "• processed_tensor: NumPy float32 array of shape (1, 224, 224, 3) with normalized values [0, 1].\n"
        "• prediction_sigmoid: Floating-point value in range [0.0, 1.0] representing tumor probability.\n"
        "• prediction_label: String ('Tumor' if prob >= 0.3 else 'No Tumor').\n"
        "• confidence_percentage: Round float percentage (0.0% to 100.0%).\n"
        "• heatmap_overlay: RGB image array combining scan and JET colormap.\n"
        "• highlight_overlay: RGB image array with red bounding box and yellow contour outline.\n"
        "• pdf_report_file: PDF document binary generated in app/static/reports/."
    )

    add_styled_heading(doc, "4.7.3 Class Diagram", level=3)
    add_body_p(
        doc,
        "System Classes & Services:\n"
        "• Config: Holds application constants (IMG_SIZE, MODEL_PATH, CLASSIFICATION_THRESHOLD).\n"
        "• ModelService: Singleton loader and prediction executor.\n"
        "• ImageService: Preprocessing, resizing, and normalization.\n"
        "• GradCAMService: Gradient extraction, JET heatmap, and contour bounding box highlight generator.\n"
        "• PDFService: ReportLab flowable PDF patient diagnostic report generator.\n"
        "• ValidationService: Extension and size checking."
    )

    add_styled_heading(doc, "4.7.4 System Activity", level=3)
    add_body_p(
        doc,
        "System Activity Flow: User Upload → Validation Guard → Preprocessor → Model Inference → Grad-CAM Engine → Tumor Contour Highlight → PDF Report Generator → HTML Template Renderer → Client Display & PDF Download."
    )

    add_styled_heading(doc, "4.8 Functional and Behavioral Modeling", level=2)

    add_styled_heading(doc, "4.8.1 Data Flow Diagram (DFD Level 0 & Level 1)", level=3)
    add_body_p(
        doc,
        "DFD Level 0 (Context Diagram): User uploads MRI file to Flask Application System, which returns Classification Label, Confidence Percentage, Heatmap Image, Tumor Highlight Image, and Downloadable PDF Diagnostic Report.\n"
        "DFD Level 1: (1.0 File Validation) → (2.0 Image Preprocessing) → (3.0 Keras Model Inference) → (4.0 Grad-CAM Heatmap & Contour Highlight Generation) → (5.0 PDF Report Microservice) → (6.0 Result Template Rendering & Download Service)."
    )

    add_styled_heading(doc, "4.8.2 Process Specification and Decision Logic", level=3)
    add_body_p(
        doc,
        "Decision Logic for Classification:\n"
        "IF prediction_probability >= 0.3 THEN label = 'Tumor', confidence = probability * 100\n"
        "ELSE label = 'No Tumor', confidence = (1 - probability) * 100"
    )

    add_styled_heading(doc, "4.9 Main Modules of New System", level=2)
    add_body_p(
        doc,
        "Main Modular Components:\n"
        "1. Web Router Module (app/routes/): Manages HTTP routes (/predict, /stats, /health, /download_report/<filename>).\n"
        "2. Model Inference Service (app/services/model_service.py): Manages Keras singleton model.\n"
        "3. Preprocessing Service (app/services/image_service.py): Prepares 224x224 RGB tensors.\n"
        "4. Explainability & Highlight Service (app/services/gradcam_service.py): Generates JET colormap overlays and red bounding box contour highlights.\n"
        "5. PDF Report Service (app/services/pdf_service.py): Programmatically builds PDF patient diagnostic reports.\n"
        "6. One-Click Launcher (start.py): Auto-creates venv, installs requirements, initializes server, and opens browser."
    )

    add_styled_heading(doc, "4.10 Selection of Hardware and Software and Justification", level=2)
    add_body_p(
        doc,
        "Selection Justification:\n"
        "• Python 3.11: Chosen for seamless compatibility with TensorFlow 2.21 and OpenCV.\n"
        "• Flask: Selected over Django for lightweight overhead and rapid RESTful rendering.\n"
        "• TensorFlow/Keras: Selected for robust CNN layers and autograd GradientTape support.\n"
        "• OpenCV: Selected for high-performance colormap blending (COLORMAP_JET) and contour detection.\n"
        "• ReportLab: Selected for programmatic, precise PDF document generation."
    )

    t_ds = doc.add_table(rows=1, cols=4)
    t_ds.alignment = WD_TABLE_ALIGNMENT.CENTER
    ds_hdr = t_ds.rows[0].cells
    ds_hdr[0].text = "Dataset Split"
    ds_hdr[1].text = "Tumor Scans"
    ds_hdr[2].text = "No Tumor Scans"
    ds_hdr[3].text = "Total Images"
    for cell in ds_hdr:
        set_cell_background(cell, "003366")
        for p in cell.paragraphs:
            for r in p.runs:
                format_run(r, size_pt=10, bold=True, color_rgb=(255, 255, 255))

    ds_data = [
        ("Training Set", "4,500", "2,000", "6,500"),
        ("Validation Set", "800", "400", "1,200"),
        ("Testing Set", "458", "208", "666"),
        ("Total Project Dataset", "5,758", "2,608", "8,366"),
    ]
    for d_split, d_tum, d_not, d_tot in ds_data:
        r_c = t_ds.add_row().cells
        r_c[0].text = d_split
        r_c[1].text = d_tum
        r_c[2].text = d_not
        r_c[3].text = d_tot
        for i, cell in enumerate(r_c):
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            for p in cell.paragraphs:
                for r in p.runs:
                    format_run(r, size_pt=10, bold=(i == 0))

    add_body_p(doc, "Table 4.1: Dataset Distribution across Train, Validation, and Test Sets")

    doc.add_page_break()

    # =============================================================
    # 5.0 SYSTEM DESIGN
    # =============================================================
    set_section_header_footer(ch_sec, header_title="5.0 System Design")
    add_styled_heading(doc, "5.0 SYSTEM DESIGN", level=1)

    add_styled_heading(doc, "5.1 Database Design and Data Structure Design", level=2)

    add_styled_heading(doc, "5.1.1 Logical Description of Data and Classes", level=3)
    add_body_p(
        doc,
        "The system uses file-system data structures and NumPy tensors rather than relational databases:\n"
        "• Model File: bestModel.keras (HDF5/Keras 3 zip format containing weights, layer configs, optimizer state).\n"
        "• Runtime Storage: app/static/uploads/ (temporary uploaded images), app/static/heatmaps/ (generated overlays & highlights), and app/static/reports/ (PDF diagnostic reports).\n"
        "• Performance Cache: In-memory python metrics dictionary rendering on /stats route."
    )

    add_styled_heading(doc, "5.2 System Procedural Design", level=2)

    add_styled_heading(doc, "5.2.1 Procedural Logic Flowchart", level=3)
    add_body_p(
        doc,
        "Procedural Logic:\n"
        "1. Receive POST request on /predict with file upload.\n"
        "2. Check if file extension is in {.png, .jpg, .jpeg}. If false, abort 400.\n"
        "3. Save file to app/static/uploads/ with unique prefix.\n"
        "4. Check file size <= 10MB. If false, remove file and abort 400.\n"
        "5. Preprocess image to 224x224x3 float32 array / 255.0.\n"
        "6. Model predicts probability p. If p >= 0.3, set label='Tumor', else 'No Tumor'.\n"
        "7. Execute Grad-CAM gradient tape pass; generate JET colormap overlay and red bounding box contour highlight image.\n"
        "8. Execute PDF Report Microservice (ReportLab) building printable patient diagnostic PDF.\n"
        "9. Render result.html template displaying prediction badge, confidence %, 3-card image grid, and 'Download PDF Diagnostic Report' button."
    )

    add_styled_heading(doc, "5.3 Input/Output and Interface Design", level=2)

    add_styled_heading(doc, "5.3.1 Samples of Forms and Interfaces", level=3)
    add_body_p(
        doc,
        "Interface Designs:\n"
        "• Homepage (index.html): Centered glassmorphic card containing drag-and-drop dropzone, file input button, 'Analyze MRI Scan' submission button, and medical disclaimer.\n"
        "• Result Page (result.html): Prediction badge (Red for Tumor, Green for No Tumor), confidence progress bar, 3-card image comparison grid (Original MRI Scan, Tumor Region Highlight, Grad-CAM Attention Map), 'Download PDF Diagnostic Report' button, and 'Analyze Another Scan' link.\n"
        "• Performance Dashboard (stats.html): Metric cards (98% Accuracy, 100% Recall, 0 False Negatives, 1.000 AUC), dataset summary table, and 4 evaluation graph plots."
    )

    add_styled_heading(doc, "5.3.2 Access Control and Security", level=3)
    add_body_p(
        doc,
        "Security Measures:\n"
        "• Input Validation: Werkzeug secure_filename() sanitizes all file paths.\n"
        "• Extension Restriction: Only .png, .jpg, and .jpeg are accepted.\n"
        "• Payload Size Limit: Flask MAX_CONTENT_LENGTH enforces 10MB maximum request size.\n"
        "• Empty Submission Guard: Frontend JS and HTML5 required attributes prevent empty form uploads.\n"
        "• Error Boundaries: Server-side try-except blocks prevent stack traces from reaching the browser."
    )

    add_styled_heading(doc, "5.4 System Architecture Design", level=2)
    add_body_p(
        doc,
        "The system follows a three-tier architecture: (1) Presentation Tier: HTML5/CSS3/Vanilla JS browser interface, (2) Application Tier: Python Flask web application server with routes (/predict, /stats, /health, /download_report/<filename>), and (3) Deep Learning & PDF Engine: TensorFlow/Keras CNN model singleton, Grad-CAM contour highlight module, and ReportLab PDF microservice."
    )

    add_styled_heading(doc, "5.5 CNN Model Architecture Details", level=2)
    add_body_p(
        doc,
        "The custom CNN architecture consists of the following sequential layers:\n"
        "• Conv2D (16 filters, 3x3 kernel, ReLU activation, input shape: 224x224x3)\n"
        "• Conv2D (32 filters, 3x3 kernel, ReLU activation)\n"
        "• MaxPool2D (2x2 pool size)\n"
        "• Conv2D (64 filters, 3x3 kernel, ReLU activation)\n"
        "• MaxPool2D (2x2 pool size)\n"
        "• Conv2D (128 filters, 3x3 kernel, ReLU activation)\n"
        "• MaxPool2D (2x2 pool size)\n"
        "• Dropout (0.25 rate)\n"
        "• Flatten Layer\n"
        "• Dense (64 units, ReLU activation)\n"
        "• Dropout (0.25 rate)\n"
        "• Dense (1 unit, Sigmoid activation head)"
    )

    doc.add_page_break()

    # =============================================================
    # 6.0 IMPLEMENTATION PLANNING AND DETAILS
    # =============================================================
    set_section_header_footer(ch_sec, header_title="6.0 Implementation Planning")
    add_styled_heading(doc, "6.0 IMPLEMENTATION PLANNING AND DETAILS", level=1)

    add_styled_heading(doc, "6.1 Implementation Environment", level=2)
    add_body_p(
        doc,
        "The system is implemented as a single-user local Flask development server and gunicorn production-ready WSGI web service. The user interface is fully GUI-based via standard modern web browsers (Chrome, Firefox, Edge)."
    )

    add_styled_heading(doc, "6.2 Program / Modules Specification", level=2)
    add_body_p(doc, "• start.py: One-click root launcher script initializing Flask server, auto-creating venv, auto-installing requirements, and auto-launching browser.")
    add_body_p(doc, "• app/app.py & app/__init__.py: Flask application factory and entrypoint.")
    add_body_p(doc, "• app/routes/: Modular blueprints for main, predict, stats, and health HTTP routes, plus PDF report download endpoint.")
    add_body_p(doc, "• app/services/: Business logic modules for model loading, preprocessing, Grad-CAM & bounding box contour highlighting, PDF report generation (pdf_service.py), validation, and stats.")
    add_body_p(doc, "• app/config.py: Configuration parameters (IMG_SIZE, MODEL_PATH, CLASSIFICATION_THRESHOLD=0.3, REPORTS_FOLDER).")

    add_styled_heading(doc, "6.3 Security Features", level=2)
    add_body_p(
        doc,
        "Security includes input file format restriction, file size enforcement (10MB limit), filename sanitization, empty submission prevention guards, and server-side exception catching rendering error.html with generic messages (HTTP 400/500)."
    )

    add_styled_heading(doc, "6.4 Coding Standards", level=2)
    add_body_p(
        doc,
        "All Python codebase modules strictly adhere to PEP8 coding standards, utilizing 4-space indentation, clear module docstrings, type annotations, and modular function separation."
    )

    add_styled_heading(doc, "6.5 Sample Coding", level=2)
    add_body_p(doc, "The code snippet below illustrates the model inference and recall-optimized threshold logic from app/services/model_service.py:")

    p_code = doc.add_paragraph()
    p_code.paragraph_format.left_indent = Inches(0.4)
    p_code.paragraph_format.space_before = Pt(6)
    p_code.paragraph_format.space_after = Pt(12)
    r_c = p_code.add_run(
        "def predict(model, processed_image):\n"
        "    \"\"\"\n"
        "    Runs model inference and applies classification threshold (0.3).\n"
        "    Returns (label: str, confidence: float 0-100).\n"
        "    \"\"\"\n"
        "    raw_pred = model.predict(processed_image, verbose=0)\n"
        "    prob = float(raw_pred[0][0])\n\n"
        "    if prob >= Config.CLASSIFICATION_THRESHOLD:  # 0.3 threshold\n"
        "        label = 'Tumor'\n"
        "        confidence = prob * 100.0\n"
        "    else:\n"
        "        label = 'No Tumor'\n"
        "        confidence = (1.0 - prob) * 100.0\n\n"
        "    return label, round(confidence, 2)\n"
    )
    format_run(r_c, font_name="Courier New", size_pt=9.5, color_rgb=(0, 51, 102))

    doc.add_page_break()
