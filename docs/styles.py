import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def set_cell_background(cell, fill_hex):
    """Sets background color of a table cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets cell padding in dxa."""
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>'
    )
    tcPr.append(tcMar)


def add_page_number(run):
    """Adds a dynamic page number field code to a paragraph run."""
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)


def format_run(run, font_name="Times New Roman", size_pt=12, bold=False, italic=False, color_rgb=(0, 0, 0)):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)


def add_styled_heading(doc, text, level=1, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True

    if level == 1:  # Chapter Heading: 16pt BOLD UPPERCASE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        format_run(run, size_pt=16, bold=True)
    elif level == 2:  # Section Heading: 14pt BOLD UPPERCASE
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.upper())
        format_run(run, size_pt=14, bold=True)
    elif level == 3:  # Subsection Heading: 12pt BOLD Title Case
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text.title())
        format_run(run, size_pt=12, bold=True)
    return p


def add_body_p(doc, text="", bold_prefix="", italic_text=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)

    if bold_prefix:
        r_prefix = p.add_run(bold_prefix)
        format_run(r_prefix, size_pt=12, bold=True)

    if text:
        r_text = p.add_run(text)
        format_run(r_text, size_pt=12, bold=False)

    if italic_text:
        r_italic = p.add_run(italic_text)
        format_run(r_italic, size_pt=12, italic=True)

    return p


def setup_section_margins(section, top=1.0, bottom=1.0, left=1.2, right=1.0):
    section.top_margin = Inches(top)
    section.bottom_margin = Inches(bottom)
    section.left_margin = Inches(left)
    section.right_margin = Inches(right)
    section.page_width = Inches(8.27)  # A4
    section.page_height = Inches(11.69)


def set_section_header_footer(section, header_title="Chapter 1.0 Introduction", is_front_matter=False):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.text = ""
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun1 = hp.add_run("23SOECE11081, 23SOECE11084\t\t")
    format_run(hrun1, size_pt=10, italic=True, color_rgb=(100, 100, 100))
    hrun2 = hp.add_run(header_title)
    format_run(hrun2, size_pt=10, bold=True, color_rgb=(100, 100, 100))

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.text = ""
    frun1 = fp.add_run("Department of Computer Engineering / IT\t\tPage ")
    format_run(frun1, size_pt=10, italic=True, color_rgb=(100, 100, 100))
    frun2 = fp.add_run()
    format_run(frun2, size_pt=10, bold=True, color_rgb=(100, 100, 100))
    add_page_number(frun2)
